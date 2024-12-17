import streamlit as st
import pandas as pd
from datetime import datetime
import time
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
from openai import AzureOpenAI
import re
import pdfplumber
import docx

# ========== Azure OpenAI Initialization ==========
client = AzureOpenAI(
    azure_endpoint="https://hkust.azure-api.net",
    api_version="2024-06-01",
    api_key="da67de83ae6f4b82985c7978a1e83c64"
)

st.set_page_config(page_title="Job Seeker Assistant (Beta)", layout="wide")

def extract_text_from_pdf(file):
    with pdfplumber.open(file) as pdf:
        text = ""
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(file):
    doc = docx.Document(file)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def local_css(file_name):
    with open(file_name) as f:
        st.markdown(f'<style>{f.read()}</style>', unsafe_allow_html=True)
        
def render_chat():
    """Render the chat messages within the chat container."""
    chat_html = '<div class="chat-container">'
    for chat_item in st.session_state["chat_history"]:
        role = chat_item["role"]
        content = chat_item["content"]
        if role == "user":
            chat_html += f'<div class="chat-message-user">👤 <b>You:</b> {content}</div>'
        else:
            chat_html += f'<div class="chat-message-ai">🤖 <b>Nexa:</b> {content}</div>'
    chat_html += '</div>'
    st.markdown(chat_html, unsafe_allow_html=True)

def render_resume_preview():
    """Render the resume preview within the resume preview box."""
    if "resume_text_display" in st.session_state:
        resume_display = st.session_state["resume_text_display"]
        resume_html = f'<div class="resume-preview">{resume_display}</div>'
        st.markdown(resume_html, unsafe_allow_html=True)

# Alternatively, embed CSS directly
def inject_css():
    """Inject custom CSS into the Streamlit app for styling."""
    css = """
    <style>
    /* General Styles */
    body {
        background-color: #f0f2f6;
    }
    /* Title Styles */
    .stTitle {
        color: #4B8BBE;
        text-align: center;
        margin-bottom: 20px;
    }
    /* File Uploader Styles */
    .file-uploader {
        background-color: #ffffff;
        padding: 2px;
        border-radius: 10px;
        box-shadow: 0 4px 8px rgba(0,0,0,0.1);
        margin-bottom: 20px;
    }
    /* Chat Container Styles */
    .chat-container {
        background-color: #f5f5f5;
        border-radius: 10px;
        padding: 50px;
        min-height: 300px;
        max-height: 600px;
        overflow-y: auto;
        box-shadow: 0 4px 8px rgba(0, 0, 0, 0.1);
        margin-bottom: 20px;
        display: flex;
        flex-direction: column;
        position: relative; /* 为了确保文字能够覆盖在容器中 */
    }
        
    .chat-container:empty::before {
        content: "Your next job is just a chat away with Nexa"; /* 提示文字 */
        color: #c8cacc; /* 暗灰色 */
        font-size: 28px; /* 字体大小 */
        position: absolute; /* 绝对定位，确保文字在容器中间 */
        top: 50%; /* 垂直居中 */
        left: 50%; /* 水平居中 */
        transform: translate(-50%, -50%); /* 精确居中 */
        font-weight: bold; /* 可选：让文字加粗 */
        white-space: nowrap;
    }
    .chat-message-user {
        background-color: #DCF8C6;
        padding: 10px 15px;
        border-radius: 15px;
        margin-bottom: 10px;
        align-self: flex-end;
        max-width: 80%;
        word-wrap: break-word;
    }
    .chat-message-ai {
        background-color: #E8E8E8;
        padding: 10px 15px;
        border-radius: 15px;
        margin-bottom: 10px;
        align-self: flex-start;
        max-width: 80%;
        word-wrap: break-word;
    }
    /* Input Box Styles */
    .stTextInput > div > div > input {
        border-radius: 10px;
        padding: 10px;
        border: 1px solid #ccc;
    }
    /* Button Styles */
   .stButton > button {
    background-color: #399bf7; /* 更柔和的浅蓝色 */
    color: #ffffff; /* 默认字体颜色白色 */
    border: none;
    border-radius: 8px;
    padding: 12px 24px;
    font-size: 16px;
    font-weight: 500;
    cursor: pointer;
    transition: background-color 0.3s, transform 0.2s, box-shadow 0.3s, color 0.3s; /* 添加字体颜色的过渡效果 */
    box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    
    /* Hover 状态 */
    .stButton > button:hover {
        background-color: #85b8f5; /* 略深的蓝色 */
        color: #ffffff; /* 保持字体颜色一致 */
        transform: translateY(-2px);
        box-shadow: 0 6px 8px rgba(0, 0, 0, 0.15);
    }
    
    /* 活动（点击）状态 */
    .stButton > button:active {
        background-color: #ffffff; /* 更深的蓝色 */
        color: #ffffff; /* 保持字体颜色一致 */
        transform: translateY(0);
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
    }
    
    /* 聚焦状态（键盘导航） */
    .stButton > button:focus {
        outline: none;
        box-shadow: 0 0 0 3px rgba(163, 201, 241, 0.5);
    }
    /* Resume Preview Box Styles */
    .resume-preview {
        background-color: #f9f9f9;
        padding: 15px;
        border: 1px solid #ddd;
        border-radius: 8px;
        max-height: 200px;
        overflow-y: auto;
        white-space: pre-wrap;
        word-wrap: break-word;
    }
    /* Scrollbar Styles */
    .chat-container::-webkit-scrollbar, .resume-preview::-webkit-scrollbar {
        width: 8px;
    }
    .chat-container::-webkit-scrollbar-thumb, .resume-preview::-webkit-scrollbar-thumb {
        background-color: rgba(0,0,0,0.2);
        border-radius: 4px;
    }
    </style>
    """
    st.markdown(css, unsafe_allow_html=True)

def get_response(message,resume_text=None):
    """
    A function to get chatbot responses from Azure OpenAI.
    It only replies to job-related keywords, ignoring unrelated queries.
    """


    job_keywords = [
        "application", "candidate", "selection", "vacancy", "resume", "cover", "letter", "portfolio", "headhunter", 
        "jobboard", "career", "careerpath", "jobsearch", "jobtitle", "employee", "recruitment", "hire", "hiring", 
        "placement", "offer", "salary", "wage", "compensation", "package", "benefits", "incentive", "perks", "reward", 
        "bonus", "raise", "pay", "hourly", "annually", "contract", "employment", "fulltime", "parttime", "temporary", 
        "freelance", "permanent", "internship", "trainee", "employmentcontract", "workpermit", "visa", "workvisa", "job", 
        "opportunity", "vacancy", "jobopening", "position", "roles", "jobtitle", "promotion", "training", "development", 
        "mentorship", "leadership", "coaching", "intern", "experience", "skillset", "advancement", "growth", "upskill", 
        "reskill", "careeradvancement", "consulting", "consultant", "trainer", "teamwork", "organization", "culture", 
        "ethics", "company", "values", "feedback", "conflict", "team", "collaboration", "communication", "time", "management", 
        "delegation", "problem", "leadership", "interpersonal", "decisions", "pressure", "growth", "challenges", "workplace", 
        "team", "organization", "consultant", "expertise", "specialist", "manager", "director", "coordinator", "executive", 
        "engineer", "developer", "scientist", "analyst", "advisor", "assistant", "recruiter", "hr", "administrator", 
        "creator", "writer", "designer", "developer", "frontend", "backend", "database", "project", "product", "ux", "ui", 
        "programmer", "software", "product manager", "business analyst", "data scientist", "data analyst", "machine", 
        "learning", "cloud", "database", "seo", "sales", "marketer", "seo", "consulting", "agency", "legal", "real", "estate", 
        "accounting", "marketing", "media", "director", "planner", "journalist", "copywriter", "editor", "photographer", 
        "scientist", "biologist", "chemist", "physicist", "actor", "producer", "receptionist", "bartender", "waiter", 
        "technician", "nurse", "teacher", "trainer", "doctor", "paramedic", "employee", "manager", "leader", "salesperson", 
        "agent", "researcher", "scientist", "developer", "strategist", "administrator", "consultant", "accountant", "broker", 
        "lawyer", "teacher", "trainer", "assistant", "driver", "operator", "contractor", "specialist", "human", "resources", 
        "feedback", "interview", "questions", "answers", "assessments", "evaluation", "assessment", "screening", "call", 
        "video", "panel", "phone", "face-to-face", "interviewing", "feedback", "assessment", "panel", "group", "reference", 
        "background", "test", "screen", "check", "employment", "tracking", "onboarding", "offboarding", "consulting", 
        "interview", "task", "project", "taskforce", "working", "remote", "flexible", "shift", "operation", "policies", 
        "solutions", "requirements", "references", "organization", "engagement", "metrics", "valuation", "review", 
        "evaluation", "annual", "quarterly", "tracking", "resume", "cv", "target", "recruitment", "headhunting", "interviews", 
        "negotiation", "compensation", "stock", "equity", "perks", "pay", "raise", "job", "salary", "benefit", "stockoptions", 
        "shares", "counseling", "advising", "tasks", "procedures", "workflow", "process", "decisionmaking", "consultation", 
        "testing", "certification", "growth", "development", "evaluation", "interviewing", "learning", "seminar", 
        "conference", "webinar", "workshops", "innovation", "digital", "marketplace", "virtual", "teamwork", "networking", 
        "social", "freelancer", "consultant", "engineer", "software", "applications", "mobile", "technology", "device", 
        "socialmedia", "startup", "design", "productivity", "system", "procedure", "project", "analyst", "developer", 
        "testing", "scrum", "agile", "productmanager", "leader", "projectmanager", "software", "solution", "development", 
        "data", "strategies", "strategy", "hiring", "leadership", "management", "digital", "design", "branding", "scaling", 
        "enterprise", "research", "corporation", "technology", "startups", "corporate", "training", "executive", 
        "director", "specialist", "expert", "mentor", "assistant", "director", "coordinator", "business", "analyst","skill","ability"
        
    ]

    lower_msg = message.lower()
    if not any(kw in lower_msg for kw in job_keywords):
        return "Sorry, I only respond to job or career-related questions."

    system_message = """
    You are a highly specialized AI assistant designed exclusively to provide expert guidance on job applications, interviews, and resume building. 
    Your purpose is to assist users in achieving their career goals through tailored advice, actionable recommendations, and detailed insights.

    Please follow these strict guidelines:

    1. **Scope of Expertise**: 
       - Focus solely on topics related to job applications, interview preparation, resume and cover letter writing, career development, and professional branding.
       - Avoid providing assistance on unrelated topics. If asked, politely explain that you are specialized in job-related guidance.

    2. **Depth and Clarity**:
       - Offer comprehensive and actionable advice. Your responses should provide step-by-step guidance or detailed examples when applicable.
       - Anticipate follow-up questions by proactively addressing common concerns or pitfalls.

    3. **Professional and Empathetic Tone**:
       - Maintain a professional, supportive, and encouraging tone. Show empathy for the challenges users face during their job search or career transitions.

    4. **Customization and Relevance**:
       - Tailor your responses to the user’s specific situation, such as their industry, level of experience, or career goals.
       - If the user provides limited context, ask clarifying questions to provide the most relevant advice.

    5. **Best Practices and Trends**:
       - Stay up-to-date with best practices in resume writing, interview techniques, and job market trends.
       - Offer insights into modern tools (e.g., ATS systems, LinkedIn optimization) and industry-specific expectations.

    6. **Proactive Suggestions**:
       - Suggest additional strategies, tips, or resources related to the user’s query. Examples include resume templates, interview question frameworks, or professional networking tips.

    7. **Specific Use Cases**:
       - **Resume and Cover Letters**: Provide tips on tailoring resumes to job descriptions, formatting for ATS systems, and crafting impactful cover letters.
       - **Interview Preparation**: Offer advice on answering behavioral questions, preparing for technical interviews, and making a strong impression.
       - **Career Guidance**: Advise on networking, personal branding, LinkedIn optimization, and identifying suitable job opportunities.

    8. **Structured Responses**:
       - Organize your responses with clear headings, bullet points, or step-by-step instructions to enhance readability and usability.
       - Where applicable, provide examples to illustrate key points (e.g., sample interview answers or resume phrases).

    Remember, your mission is to empower users with the knowledge, tools, and confidence to excel in their job search and career advancement. Keep your focus sharp and responses actionable.
    """
    if resume_text:
        system_message += f"\n\nUser's Resume:\n{resume_text}\n\nPlease consider this information when providing your response."

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        temperature=1,
        messages = [
            {"role": "system", "content": system_message},
            {"role": "user", "content": message}
        ]
    )
    return response.choices[0].message.content




# ========== Streamlit Main App ==========

@st.cache_data
def your_function_name():
    
    pass
    
def load_data():
    """
    Loads the English-version Excel data for job listings.
    Make sure the file path and column names match your dataset structure.
    """
    file_path = 'Recruitment_Data_English 1.xlsx'
    
    return pd.read_excel(file_path, engine='openpyxl')

import os
from docx import Document

def load_docx(file_path):
    """
    Load content from a .docx file and return its text.
    """
    # 检查文件是否存在
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    # 确保文件扩展名正确
    if not file_path.endswith('.docx'):
        raise ValueError("Invalid file type. Please provide a .docx file.")

    # 读取 .docx 文件内容
    try:
        doc = Document(file_path)
        content = "\n".join([para.text for para in doc.paragraphs])
        return content
    except Exception as e:
        raise ValueError(f"Error reading the DOCX file: {e}")

# 示例调用
file_path = "Recruitment_Data_English 1.xlsx"  # 确保此文件存在
try:
    text = load_docx(file_path)
    print("Document content loaded successfully:")
    print(text)
except Exception as e:
    print(e)


@st.cache_resource
def load_embedding_model(model_name='sentence-transformers/all-MiniLM-L6-v2'):
    """
    Loads the sentence transformer model for encoding text.
    (Switched to an English-compatible model for demonstration.)
    """
    return SentenceTransformer(model_name)

@st.cache_data
def compute_embeddings(df, _model):
    """
    Combines certain columns (Working city + Available positions + Company description)
    into a single string and encodes them with the Transformer model for similarity search.
    """
    df['combined_info'] = (
        df['Working city'].astype(str) + ' ' +
        df['Available positions'].astype(str) + ' ' +
        df['Company description'].astype(str)
    )
    embeddings = _model.encode(df['combined_info'].tolist(), show_progress_bar=True)
    return embeddings


def display_job_info(df, page_num, items_per_page=10):
    """
    Displays job information in a paginated format with an additional check for email links.
    """
    total_items = len(df)
    if total_items == 0:
        st.info("No data available.")
        return

    total_pages = (total_items - 1) // items_per_page + 1
    if page_num < 1:
        page_num = 1
    elif page_num > total_pages:
        page_num = total_pages
        st.warning(f"Out of range. Jumping to last page: Page {total_pages}")

    start_idx = (page_num - 1) * items_per_page
    end_idx = start_idx + items_per_page
    subset = df.iloc[start_idx:end_idx]

    if subset.empty:
        st.warning("No results on this page, try another page number.")
        return

    # Pull the user_position_input from session_state
    user_position_input = st.session_state.get("user_position_input", "").strip().lower()

    # --- HTML/CSS for job cards ---
    st.markdown(""" 
    <head>
        <link href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/5.15.4/css/all.min.css" rel="stylesheet">
    </head>
    <style>
    .job-card {
        border: 1px solid #c4c4c4;
        border-radius: 12px; /* 增加圆角半径 */
        padding: 2px; /* 增加内边距 */
        margin-bottom: 5px; /* 增加底部间距 */
        background-color: #ffffff;
        display: flex;
        flex-direction: row;
        justify-content: space-between;
        align-items: flex-start; 
        transition: transform 0.2s, box-shadow 0.2s; /* 添加过渡效果 */
    }
    .job-card:hover {
        transform: translateY(-5px); /* 悬停时稍微上移 */
        box-shadow: 0 8px 16px rgba(0, 0, 0, 0.2); /* 悬停时加深阴影 */
    }
    .job-left {
        width: 75%;
    }
    .job-title {
        font-size: 2.0rem; /* 调整字体大小 */
        font-weight: 700;
        margin-bottom: 10px;
        color: #333;
    }
    .job-meta {
        font-size: 1rem;
        color: #555;
        margin-bottom: 12px;
    }
    .job-meta span { 
        margin-right: 20px; 
    }
    .company-tag {
        display: inline-block;
        background-color: #e3f2fd;
        color: #1976d2;
        border-radius: 8px;
        padding: 3px 8px;
        margin-right: 5px;
        font-size: 0.85rem;
        border: 1px solid #1976d2;
    }
    .deadline-text {
        color: #d9534f;
    }
    .job-positions {
        margin-top: 10px;
        color: #333;
    }
    .pos-badge {
        display: inline-block;
        border-radius: 6px;
        padding: 5px 12px;
        margin: 4px 6px 0 0;
        font-size: 0.9rem;
        font-weight: 500;
    }
    .pos-green {
        border: 2px solid #4caf50;
        color: #2e7d32;
        background-color: #e8f5e9; /* 添加背景颜色 */
    }
    .pos-gray {
        border: 2px solid #a8a5a5;
        color: #a8a5a5;
        background-color: #f5f7f7; /* 添加背景颜色 */
    }
    .job-right {
        width: 23%;
        display: flex;
        flex-direction: column;
        align-items: flex-end;
        gap: 0px;
        margin-top: 5px;
    }
    .button-container {
        display: flex;
        justify-content: flex-start;
        gap: 10px;
    }
   a.styled-btn-detail {
        display: inline-block;
        text-decoration: none;
        background-color: #ffffff; /* 白色背景 */
        color: #333333; /* 深灰色字体 */
        font-size: 1.2rem;
        padding: 8px 15px; /* 增加内边距 */
        border-radius: 8px;
        border: 2px solid #a8a5a5; /* 灰色边框 */
        transition: background-color 0.3s ease, color 0.3s ease, transform 0.2s;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1); /* 添加轻微阴影 */
        font-weight: 600;
        margin-right: 20px;
    }
    a.styled-btn-detail:hover {
        background-color: #f0f0f0; /* 悬停时背景稍微变深 */
        color: #000000; /* 悬停时字体变为黑色 */
        transform: translateY(-2px); /* 悬停时稍微上移 */
        box-shadow: 0 4px 8px rgba(0, 0, 0, 0.15); /* 增强阴影 */
    }
    a.styled-btn-detail:active {
        background-color: #e0e0e0; /* 点击时背景进一步变深 */
        color: #000000;
        transform: translateY(0); /* 点击后恢复位置 */
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
    }
    a.styled-btn-apply {
        display: inline-block;
        text-decoration: none;
        background-color: #399bf7; /* 蓝色背景 */
        color: #ffffff; /* 白色字体 */
        font-size: 1.2rem;
        padding: 8px 15px; /* 增加内边距 */
        border-radius: 8px;
        border: 2px solid #399bf7; /* 蓝色边框，与背景相同 */
        transition: background-color 0.3s ease, color 0.3s ease, transform 0.2s;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1); /* 添加轻微阴影 */
        font-weight: 600;
    }
    a.styled-btn-apply:hover {
        background-color: #1565c0; /* 悬停时背景变深 */
        color: #ffffff;
        transform: translateY(-2px);
        box-shadow: 0 4px 8px rgba(0, 0, 0, 0.15); /* 增强阴影 */
    }
    a.styled-btn-apply:active {
        background-color: #0d47a1; /* 点击时背景进一步变深 */
        color: #ffffff;
        transform: translateY(0);
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
    }
    .company-description {
        display: inline-flex;
        align-items: center;
        background-color: #fcf8d4;  
        color: #333;  
        border-radius: 8px;  
        padding: 6px 12px;  
        font-size: 0.85rem;  
        margin-top: 10px; 
        margin-bottom: 0px;
    }
    .company-description i {
        margin-right: 6px;  
    }
    .warning-text {
        color: #d9534f;
        font-weight: bold;
        margin-top: 10px;
    }
    </style>
    """, unsafe_allow_html=True)




    current_date = datetime.now().date()

    for idx, row in subset.iterrows():
        detail_link = row.get('Detail link', '#')
        apply_link = row.get('Application link', '#')

        # Check if the application link is an email address
        email_pattern = r"^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$"
        is_email = re.match(email_pattern, apply_link) is not None

        tags = row['Company label']
        if isinstance(tags, list):
            tags_html = "".join(f'<span class="company-tag">{t}</span>' for t in tags)
        else:
            tags_html = f"<span class='company-tag'>{tags}</span>" if tags else ""

        # Deadline check
        expired = False
        if pd.notnull(row['Deadline']):
            if row['Deadline'].date() < current_date:
                expired = True
            deadline_str = row['Deadline'].date().strftime("%Y-%m-%d")
        else:
            deadline_str = "No Info"

        if expired:
            date_html = f'<span style="color:red;"> <i class="fas fa-clock"></i> {deadline_str} (Expired)</span>'
        else:
            date_html = f'<span style="color:black;"> <i class="fas fa-clock"></i> {deadline_str}</span>'

        with st.container():
            st.markdown('<div class="job-card">', unsafe_allow_html=True)
            
            # 1) Company Name
            st.markdown(f'<div class="job-title">{row["Company Name"]}</div>', unsafe_allow_html=True)
            # 2) City / Tags / Deadline
            st.markdown(
                f'<div class="job-meta">'
                f'<span><i class="fas fa-map-marker-alt"></i> {row["Working city"]}</span>'
                f'<span><i class="fas fa-tags"></i> {tags_html}</span>'
                f'{date_html}'
                f'</div>',
                unsafe_allow_html=True
            )

            # 3) Available positions
            positions_raw = str(row["Available positions"]) if row["Available positions"] else ""
            positions = re.split(r'[-]+', positions_raw.strip())

            pos_html_list = []
            for p in positions:
                p_clean = p.strip()
                if not p_clean:
                    continue
                if user_position_input and (user_position_input in p_clean.lower()):
                    pos_html_list.append(f'<span class="pos-badge pos-green">{p_clean}</span>')
                else:
                    pos_html_list.append(f'<span class="pos-badge pos-gray">{p_clean}</span>')

            pos_html = " ".join(pos_html_list)
            st.markdown(f'<div class="job-positions">{pos_html}</div>', unsafe_allow_html=True)

            # --- Company Description Section as a Tag ---
            company_description = str(row.get('Company description', 'No description available.'))
            st.markdown(f'<span class="company-description"><i class="fas fa-info-circle"></i>{company_description}</span>', unsafe_allow_html=True)

            st.markdown('</div>', unsafe_allow_html=True)  # close job-left


            if is_email:
                st.markdown(f"""
                    <div class="button-container">
                        <a class="styled-btn-detail" href="{detail_link}" target="_blank">Detail Link</a>
                        <a class="styled-btn-apply" href="mailto:{apply_link}">Application Email</a>
                    </div>
                """, unsafe_allow_html=True)
                
                st.markdown('</div>', unsafe_allow_html=True)
                # 使用 st.warning 显示警告消息
                st.warning(
    f"Please send your resume to [{apply_link}](mailto:{apply_link}) for application. More details, please check the [Detail Link]({detail_link}). Thank you!"
)

            else:
                st.markdown(f"""
                    <div class="button-container">
                        <a class="styled-btn-detail" href="{detail_link}" target="_blank">Detail Link</a>
                        <a class="styled-btn-apply" href="{apply_link}" target="_blank">Application Link</a>
                    </div>
                """, unsafe_allow_html=True)

            st.markdown('</div>', unsafe_allow_html=True)  # close job-card

    return page_num, total_pages




def main():
    
    # App layout
    col1, col2 = st.columns([3,2], gap="large")
    inject_css()

    with col1:
        st.title("🌟Joblytic (Beta)")
        st.subheader("Your entire job search. Powered by one App")

        with st.spinner('Loading recruitment data, please wait...'):
            data = load_data()
            time.sleep(1)

        global model
        model = load_embedding_model()

        if 'exact_matches' not in st.session_state:
            st.session_state['exact_matches'] = pd.DataFrame()
        if 'similar_matches' not in st.session_state:
            st.session_state['similar_matches'] = pd.DataFrame()

        # --- Parse Company Labels (Fix: only split by comma) ---
        # If a row has "Media/live broadcast, state-owned enterprises", it means two tags:
        #   ["Media/live broadcast", "state-owned enterprises"].
        all_tags = set()
        parsed_tags = []
        for idx, row in data.iterrows():
            raw = str(row.get('Company label', '')).strip()
            if raw and raw.lower() != 'nan':
                # Split ONLY by comma, so "Media/live broadcast" remains a single tag
                tag_list = [t.strip() for t in raw.split(',') if t.strip()]
            else:
                tag_list = []
            parsed_tags.append(tag_list)
            all_tags.update(tag_list)
        data['Company label'] = parsed_tags
        all_tags = sorted(list(all_tags))

        with st.expander("🔍 Search Filters", expanded=True):
            city_input = st.text_input("City:", "Shenzhen")
            position_input = st.text_input("Position:", "Data analysis")
            st.session_state["user_position_input"] = position_input  # store user input (lowercase match used later)

            selected_tags = st.multiselect("Filter by Company Label (Multiple Select)", options=all_tags)
            
            hide_expired_jobs = st.toggle("Hide jobs with expired deadlines", value=True)

            if st.button("Search Jobs"):
                if selected_tags:
                    # If any selected tag appears in the company's label list, we keep it
                    mask = data['Company label'].apply(lambda tags: any(t in tags for t in selected_tags))
                    filtered_data = data[mask].copy()
                else:
                    filtered_data = data.copy()

                data_embeddings = compute_embeddings(filtered_data, model)
                user_query = f"{city_input} {position_input}"
                query_embedding = model.encode([user_query.lower()], show_progress_bar=False)
                similarity_scores = cosine_similarity(query_embedding, data_embeddings).flatten()
                filtered_data['similarity'] = similarity_scores

                # Make city & position match case-insensitive
                city_lower = city_input.lower().strip()
                position_lower = position_input.lower().strip()

                filtered_data['city_match'] = filtered_data['Working city'].astype(str).apply(
                    lambda x: city_lower in x.lower()
                )
                filtered_data['position_match'] = filtered_data['Available positions'].astype(str).apply(
                    lambda x: position_lower in x.lower()
                )
                filtered_data['city_position_match'] = filtered_data['city_match'] & filtered_data['position_match']

                # Adjust similarity
                # Adjust similarity
                def compute_adjusted_similarity(row):
                    base_sim = row['similarity']
                    if row['city_position_match']:
                        return base_sim * 1.2
                    elif row['city_match'] or row['position_match']:
                        return base_sim * 1.1
                    else:
                        return base_sim
                
                filtered_data['adjusted_similarity'] = filtered_data.apply(compute_adjusted_similarity, axis=1)
                
                # Deadline effect
                filtered_data['Deadline'] = pd.to_datetime(filtered_data['Deadline'], errors='coerce')
                current_date = datetime.now().date()
                
                # Assign priority based on deadline status
                def get_deadline_priority(row):
                    if pd.isnull(row['Deadline']):
                        return 2  # No Info
                    elif row['Deadline'].date() >= current_date:
                        return 1  # Not expired
                    else:
                        return 3  # Expired
                
                filtered_data['deadline_priority'] = filtered_data.apply(get_deadline_priority, axis=1)
                
                # Adjust similarity for expired jobs
                filtered_data['adjusted_similarity'] = filtered_data.apply(
                    lambda row: row['adjusted_similarity'] * 0.5 if row['deadline_priority'] == 3 else row['adjusted_similarity'],
                    axis=1
                )
                
                # Filter out expired jobs if checkbox is selected
                if hide_expired_jobs:
                    filtered_data = filtered_data[filtered_data['deadline_priority'] != 3]
                
                # Sort by priority and similarity
                filtered_data = filtered_data.sort_values(by=['deadline_priority', 'adjusted_similarity'], ascending=[True, False])
                
                # Separate exact and similar matches
                exact_matches = filtered_data[filtered_data['city_position_match']].copy()
                similar_matches = filtered_data[~filtered_data['city_position_match']].copy()
                
                # Update session state
                st.session_state['exact_matches'] = exact_matches
                st.session_state['similar_matches'] = similar_matches



        exact_matches = st.session_state['exact_matches']
        similar_matches = st.session_state['similar_matches']

        if exact_matches.empty and similar_matches.empty:
            st.info("Please set your filters and click 'Search Jobs'.")
        else:
            tab1, tab2 = st.tabs(["Exact Match", "Similar Match"])
            with tab1:
                if exact_matches.empty:
                    st.warning("No exact matches found, but more jobs opportunities are on the way!")
                else:
                    st.subheader("I found those jobs that fit your needs:")
                    exact_total = len(exact_matches)
                    exact_pages = (exact_total - 1) // 10 + 1
                    exact_page_num = st.number_input(
                        f"Exact matches, total {exact_total}. Choose page:",
                        min_value=1,
                        max_value=exact_pages,
                        value=1,
                        step=1,
                        key="exact_matches_page"
                    )
                    display_job_info(exact_matches, exact_page_num, items_per_page=10)

            with tab2:
                st.subheader("Jobs opportunities you might also interested in(Top 20):")
                similar_top20 = similar_matches.head(20)
                if similar_top20.empty:
                    st.write("No similar job recommendations.")
                else:
                    st.write(f"Total {len(similar_matches)} similar matches. Showing top 20 below:")
                    display_job_info(similar_top20, 1, items_per_page=20)

    # ========== Right Column: AI Chatbox ==========
    with col2:
        st.markdown("<h1 style='text-align:center;'>🤖 Nexa: Your AI Career Consultant</h1>", unsafe_allow_html=True)
        
        # File Upload Section
        st.subheader("📄 Upload Your Resume")
        uploaded_file = st.file_uploader("Upload your resume (PDF or DOCX)", type=["pdf", "docx"])
        st.markdown('</div>', unsafe_allow_html=True)
        
        resume_text = None
        if uploaded_file is not None:
            try:
                if uploaded_file.type == "application/pdf":
                    resume_text = extract_text_from_pdf(uploaded_file)
                elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    resume_text = extract_text_from_docx(uploaded_file)
                st.success("✅ Upload successful!")
                st.markdown("### 📄 Resume Preview:")
                # Limit the resume preview to first 1000 characters
                if len(resume_text) > 3000:
                    st.session_state["resume_text_display"] = resume_text[:3000] + "..."
                else:
                    st.session_state["resume_text_display"] = resume_text
                render_resume_preview()
                st.markdown('<div class="file-uploader">', unsafe_allow_html=True)
            except Exception as e:
                st.error(f"❌ Upload failed, please check your file: {e}")
        
        st.markdown("### 💬 Chat with Nexa")

        # Initialize chat history
        if "chat_history" not in st.session_state:
            st.session_state["chat_history"] = []
        if "chat_input" not in st.session_state:
            st.session_state["chat_input"] = ""
        
        # Chat Container
        render_chat()
        
        # Chat Input
        user_input = st.text_input(
            "Type your message here...",
            value=st.session_state["chat_input"],
            placeholder="Ask me anything about your career..."
        )
        
        # Buttons: Send and Clear Chat
        col_send, col_clear = st.columns([1, 1], gap="small")
        
        with col_send:
            if st.button("Send"):
                if user_input.strip():
                    st.session_state["chat_history"].append({"role": "user", "content": user_input})
                    reply = get_response(user_input, resume_text=resume_text)
                    st.session_state["chat_history"].append({"role": "assistant", "content": reply})
                    st.session_state["chat_input"] = ""
                    st.rerun()  # Replace st.experimental_rerun() with st.rerun()
                else:
                    st.warning("Please enter a message before sending.")
        
        with col_clear:
            if st.button("Clear Chat History"):
                st.session_state["chat_history"] = []
                st.session_state["chat_input"] = ""
                st.rerun()  # Replace st.experimental_rerun() with st.rerun()

# ========= 4. Run the Application ==========
if __name__ == "__main__":
    main()
